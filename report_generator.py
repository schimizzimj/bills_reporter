# Imports
import sys
import datetime
import docx
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE

# Constants
CALIBRI = 'Calibri'
STATES = {
    'Alabama': 'AL',
    'Alaska': 'AK',
    'American Samoa': 'AS',
    'Arizona': 'AZ',
    'Arkansas': 'AR',
    'California': 'CA',
    'Colorado': 'CO',
    'Connecticut': 'CT',
    'Delaware': 'DE',
    'District of Columbia': 'DC',
    'Florida': 'FL',
    'Georgia': 'GA',
    'Guam': 'GU',
    'Hawaii': 'HI',
    'Idaho': 'ID',
    'Illinois': 'IL',
    'Indiana': 'IN',
    'Iowa': 'IA',
    'Kansas': 'KS',
    'Kentucky': 'KY',
    'Louisiana': 'LA',
    'Maine': 'ME',
    'Maryland': 'MD',
    'Massachusetts': 'MA',
    'Michigan': 'MI',
    'Minnesota': 'MN',
    'Mississippi': 'MS',
    'Missouri': 'MO',
    'Montana': 'MT',
    'Nebraska': 'NE',
    'Nevada': 'NV',
    'New Hampshire': 'NH',
    'New Jersey': 'NJ',
    'New Mexico': 'NM',
    'New York': 'NY',
    'North Carolina': 'NC',
    'North Dakota': 'ND',
    'Northern Mariana Islands':'MP',
    'Ohio': 'OH',
    'Oklahoma': 'OK',
    'Oregon': 'OR',
    'Pennsylvania': 'PA',
    'Puerto Rico': 'PR',
    'Rhode Island': 'RI',
    'South Carolina': 'SC',
    'South Dakota': 'SD',
    'Tennessee': 'TN',
    'Texas': 'TX',
    'Utah': 'UT',
    'Vermont': 'VT',
    'Virgin Islands': 'VI',
    'Virginia': 'VA',
    'Washington': 'WA',
    'West Virginia': 'WV',
    'Wisconsin': 'WI',
    'Wyoming': 'WY'
}

def get_bill_link(state, bill_number):
    state_code = STATES[state]
    return f"https://legiscan.com/{state_code}/bill/{bill_number}/2021"

def read_bills_file(filename):
    # Read in export file
    bills = []
    headers = []
    with open(filename, "r", encoding='windows-1252') as f:
        for index, line in enumerate(f):
            line = line.strip().split("\t")
            if index == 0:
                headers = line
            else:    
                bill = {}
                for i, field in enumerate(line):
                    bill[headers[i]] = field
                bills.append(bill)
    return bills

def create_word_doc(bills, output_name='report.docx'):
    # Create output document
    document = Document()
    section = document.sections[0]
    header = section.header


    # Define needed styles
    styles = document.styles

    base_style = styles.add_style('Base Style', WD_STYLE_TYPE.CHARACTER)
    base_style.font.name = CALIBRI
    base_style.font.italic = False

    italic_base_style = styles.add_style('Italic Base Style', WD_STYLE_TYPE.PARAGRAPH)
    italic_base_style.font.name = CALIBRI
    italic_base_style.font.italic = True

    state_style = styles.add_style('State Style', WD_STYLE_TYPE.PARAGRAPH)
    state_style.font.name = CALIBRI
    state_style.font.bold = True
    state_style.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    bill_code_style = styles.add_style('Bill Code Style', WD_STYLE_TYPE.PARAGRAPH)
    bill_code_style.font.name = CALIBRI
    bill_code_style.font.size = Pt(14)
    bill_code_style.font.bold = True

    bill_name_style = styles.add_style('Bill Name Style', WD_STYLE_TYPE.CHARACTER)
    bill_name_style.font.name = CALIBRI
    bill_name_style.font.size = Pt(14)
    bill_name_style.font.bold = False

    description_style = styles.add_style('Description Style', WD_STYLE_TYPE.PARAGRAPH)

    link_style = styles.add_style('Link Style', WD_STYLE_TYPE.PARAGRAPH)
    link_style.font.name = CALIBRI
    link_style.font.underline = True
    link_style.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)


    # Add timestamp for header
    header_style = document.styles['Header']
    header_style.font.name = CALIBRI
    now = datetime.datetime.now().astimezone().strftime("%m/%d/%Y %I:%M:%S%p %Z")
    header.paragraphs[0].text = now


    # Add paragraph for each bill
    for bill in bills:
        document.add_paragraph(bill['Legislature'], style=state_style)
        document.add_paragraph(bill['Bill Number'], style=bill_code_style).add_run(' ' + bill['Shortened Title'], style=bill_name_style)
        document.add_paragraph(bill['Description'], style=description_style)

        link = document.add_paragraph(style=link_style)
        part = link.part
        r_id = part.relate_to(get_bill_link(bill['Legislature'], bill['Bill Number']), docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        new_run.append(rPr)
        new_run.text = 'Bill Text'
        hyperlink.append(new_run)
        link._p.append(hyperlink)

        document.add_paragraph('Sponsor: ', style=italic_base_style).add_run(bill['Sponsor'], style=base_style)
        document.add_paragraph('Latest Action: ', style=italic_base_style).add_run(bill['Latest Action'], style=base_style)

        document.add_paragraph()
        document.add_paragraph()

    document.save(output_name)

def main(input_file):
    # Read in file to dict
    bills = read_bills_file(input_file)

    # Create formatted word document
    create_word_doc(bills)

if __name__ == '__main__':
    main(sys.argv[1])